from docx import Document

doc = Document("E:\\Manas.docx")
par= len(doc.paragraphs)
for i in range(par):
    print(doc.paragraphs[i].text)